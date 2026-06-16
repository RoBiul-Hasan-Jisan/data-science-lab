from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# Create document
doc = Document()

# --- HEADER ---
header_table = doc.add_table(rows=1, cols=2)
header_table.autofit = False
header_table.columns[0].width = Inches(4.5)
header_table.columns[1].width = Inches(2.5)

# Left side of header (Name & title)
cell_left = header_table.cell(0, 0)
cell_left.text = "Ayesha Shiddika Jasika\nMatrimonial Bio-data"
cell_left.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
run_left = cell_left.paragraphs[0].runs[0]
run_left.font.size = Pt(20)
run_left.font.bold = True
run_left.font.color.rgb = RGBColor(255, 255, 255)

# Right side of header (Photo placeholder)
cell_right = header_table.cell(0, 1)
cell_right.text = "[Insert Photo Here]"
cell_right.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run_right = cell_right.paragraphs[0].runs[0]
run_right.font.size = Pt(12)
run_right.font.color.rgb = RGBColor(255, 255, 255)

# Style header background
for row in header_table.rows:
    for cell in row.cells:
        shading_elm = parse_xml(r'<w:shd {} w:fill="2F5496"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading_elm)

doc.add_paragraph()

# Function for section heading
def add_section(title):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = title
    cell.paragraphs[0].runs[0].font.size = Pt(14)
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    shading_elm = parse_xml(r'<w:shd {} w:fill="4472C4"/>'.format(nsdecls('w')))
    cell._tc.get_or_add_tcPr().append(shading_elm)

# Function for detail row
def add_detail(label, value):
    t = doc.add_table(rows=1, cols=2)
    t.autofit = True
    t.cell(0, 0).text = label
    t.cell(0, 0).paragraphs[0].runs[0].font.bold = True
    t.cell(0, 1).text = value

# --- Personal Information ---
add_section("Personal Information")
personal_info = [
    ("Date of Birth", "30 May 2004"),
    ("Age", "21 years (as of 2025)"),
    ("Blood Group", "A"),
    ("Marital Status", "Unmarried"),
    ("Nationality", "Bangladeshi"),
    ("Religion", "Islam")
]
for label, value in personal_info:
    add_detail(label, value)

doc.add_paragraph()

# --- Education ---
add_section("Educational Qualifications")
doc.add_paragraph("1. Honours (Ongoing)\n   - Cox’s Bazar Government College\n   - Department: Political Science\n   - Session: 2023–2024")
doc.add_paragraph("2. Higher Secondary Certificate (HSC) (Ongoing)\n   - Cox’s Bazar Government Women’s College\n   - Group: Humanities\n   - Batch: 2025")

doc.add_paragraph()

# --- Family Information ---
add_section("Family Information")
add_detail("Father’s Name", "Abul Kalam Azad (Businessman)")
add_detail("Mother’s Name", "Jousna Akter (Housewife)")
doc.add_paragraph("Siblings:\n1. Abu Shuhan Ayuth – SSC Candidate (Batch-2025)\n2. Ashpia Jannat Lobaba – Class 7\n3. Wasika Jannat Subah – Class 1")

doc.add_paragraph()

# --- Address ---
add_section("Permanent Address")
doc.add_paragraph("Bacho Munshir Paro, Amirabad, Chittagong, Bangladesh")

# --- Contact ---
add_section("Contact Information")
add_detail("Mobile", "[Insert]")
add_detail("Email", "[Insert]")

# --- Declaration ---
add_section("Declaration")
doc.add_paragraph("I hereby declare that the above information is true and accurate to the best of my knowledge.\n\nSignature: _____________________\nDate: __________________________")

# Save
file_path = 'Ayesha_Shiddika_Jasika_Advanced_Matrimonial_CV.docx'


doc.save(file_path)


file_path
